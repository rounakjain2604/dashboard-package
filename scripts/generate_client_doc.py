"""
Generate a professional, print-ready DOCX document for client delivery.
IB-Grade DCF Valuation Engine — Client Package Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Constants ───────────────────────────────────────────────────────────
FONT_NAME = "Verdana"
COLOR_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)     # Dark navy
COLOR_ACCENT = RGBColor(0x2E, 0x86, 0xAB)       # Teal blue
COLOR_DARK = RGBColor(0x33, 0x33, 0x33)          # Near-black text
COLOR_MEDIUM = RGBColor(0x55, 0x55, 0x55)        # Gray body
COLOR_LIGHT_BG = RGBColor(0xF0, 0xF4, 0xF8)     # Light blue-gray bg
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_HEADER = RGBColor(0x1B, 0x3A, 0x5C)  # Navy header
COLOR_TABLE_ALT = RGBColor(0xF7, 0xF9, 0xFB)     # Alternating row
COLOR_GREEN = RGBColor(0x27, 0xAE, 0x60)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "IB_Grade_DCF_Engine_Client_Document.docx")


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_font(run, size=9, bold=False, color=COLOR_DARK, italic=False, name=FONT_NAME):
    """Apply consistent font styling to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    # Set East-Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}" w:cs="{name}"/>')
        rPr.insert(0, rFonts)


def add_heading_styled(doc, text, level=1):
    """Add a styled heading with Verdana font."""
    p = doc.add_paragraph()
    if level == 1:
        p.space_before = Pt(16)
        p.space_after = Pt(6)
        run = p.add_run(text.upper())
        set_font(run, size=14, bold=True, color=COLOR_PRIMARY)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="2E86AB"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        p.space_before = Pt(12)
        p.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run, size=11, bold=True, color=COLOR_ACCENT)
    elif level == 3:
        p.space_before = Pt(8)
        p.space_after = Pt(3)
        run = p.add_run(text)
        set_font(run, size=10, bold=True, color=COLOR_PRIMARY)
    return p


def add_body(doc, text, space_after=4, space_before=0):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    p.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    set_font(run, size=9, color=COLOR_DARK)
    return p


def add_bullet(doc, text, indent_level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    p.space_before = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    left_indent = Cm(1.2 + indent_level * 0.8)
    p.paragraph_format.left_indent = left_indent
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run("•  " + text)
    set_font(run, size=8.5, color=COLOR_DARK)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a professionally formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.space_before = Pt(3)
        p.space_after = Pt(3)
        run = p.add_run(header)
        set_font(run, size=8, bold=True, color=COLOR_WHITE)
        set_cell_shading(cell, "1B3A5C")

    # Style data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.space_before = Pt(2)
            p.space_after = Pt(2)
            run = p.add_run(str(val))
            set_font(run, size=8, color=COLOR_DARK)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F7F9FB")

    # Apply column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph().space_after = Pt(4)
    return table


def add_kpi_box(doc, items):
    """Add a row of KPI-style highlight boxes using a table."""
    table = doc.add_table(rows=2, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(items):
        # Value cell
        cell_top = table.rows[0].cells[i]
        cell_top.text = ""
        p = cell_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(6)
        p.space_after = Pt(2)
        run = p.add_run(value)
        set_font(run, size=16, bold=True, color=COLOR_ACCENT)
        set_cell_shading(cell_top, "F0F4F8")
        # Label cell
        cell_bot = table.rows[1].cells[i]
        cell_bot.text = ""
        p2 = cell_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.space_before = Pt(1)
        p2.space_after = Pt(6)
        run2 = p2.add_run(label)
        set_font(run2, size=7.5, bold=True, color=COLOR_PRIMARY)
        set_cell_shading(cell_bot, "F0F4F8")

    # Remove borders for KPI box
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph().space_after = Pt(4)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.oxml import OxmlElement
    br_elem = OxmlElement('w:br')
    br_elem.set(qn('w:type'), 'page')
    run._element.append(br_elem)


# ═════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # ─── Page setup ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)

    # ─── Set default font ────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(9)
    font.color.rgb = COLOR_DARK

    # ═════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ═════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph().space_after = Pt(6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IB-Grade DCF Valuation Engine")
    set_font(run, size=26, bold=True, color=COLOR_PRIMARY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(6)
    run2 = p2.add_run("VERSION 13.0")
    set_font(run2, size=14, bold=True, color=COLOR_ACCENT)

    # Decorative line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(20)
    run3 = p3.add_run("━" * 50)
    set_font(run3, size=10, color=COLOR_ACCENT)

    # Subtitle
    subtitle_lines = [
        "A Fully Automated, Production-Ready DCF Modelling Platform",
        "Interactive Dashboard  ·  Formula-Linked Excel  ·  PDF Memo",
        "",
        "Client Product Guide & Technical Walkthrough",
    ]
    for line in subtitle_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(4)
        run = p.add_run(line)
        if line == "":
            continue
        if "Client Product" in line:
            set_font(run, size=11, bold=True, color=COLOR_PRIMARY)
        else:
            set_font(run, size=10, color=COLOR_MEDIUM)

    for _ in range(4):
        doc.add_paragraph().space_after = Pt(6)

    # Confidentiality notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — FOR INTENDED RECIPIENT ONLY")
    set_font(run, size=8, bold=True, color=COLOR_MEDIUM, italic=True)

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)

    toc_items = [
        ("1.", "Executive Summary"),
        ("2.", "What You Are Getting"),
        ("3.", "Unique Selling Points (USP)"),
        ("4.", "Platform Walkthrough"),
        ("5.", "Dashboard — 10-Tab Breakdown"),
        ("6.", "Excel Export — 17-Tab Workbook"),
        ("7.", "Complete File & Folder Reference"),
        ("8.", "Installation & Setup Guide"),
        ("9.", "How to Load Your Own Data"),
        ("10.", "Configuration Reference"),
        ("11.", "Deployment Options"),
        ("12.", "Extension & Customisation Guide"),
        ("13.", "API Reference"),
        ("14.", "Version History & Quality Assurance"),
        ("15.", "Support & Next Steps"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(1.2))
        run_num = p.add_run(num)
        set_font(run_num, size=9, bold=True, color=COLOR_ACCENT)
        run_title = p.add_run("  " + title)
        set_font(run_title, size=9, color=COLOR_DARK)

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Executive Summary", level=1)

    add_body(doc,
        "This document presents the IB-Grade DCF Valuation Engine — a fully automated, "
        "production-ready platform that transforms raw financial data and user assumptions "
        "into institutional-quality valuation outputs. The engine is designed for analysts, "
        "investors, buyers, and advisory teams who need fast, auditable, and repeatable "
        "financial modelling without manual spreadsheet work."
    )

    add_body(doc,
        "The platform runs a 16-step financial modelling pipeline entirely in Python, "
        "producing three distinct deliverables from one set of inputs:"
    )

    add_kpi_box(doc, [
        ("INTERACTIVE DASHBOARD", "10 Tabs"),
        ("FORMULA-LINKED EXCEL", "17 Tabs"),
        ("PDF INVESTMENT MEMO", "10-15 Pages"),
    ])

    add_body(doc,
        "Every output is generated from the same underlying model logic — the dashboard "
        "shows live analytics, the Excel workbook contains real formulas (not pasted values) "
        "that recalculate when you change assumptions, and the PDF memo is ready to attach "
        "to a deal package."
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  2. WHAT YOU ARE GETTING
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. What You Are Getting", level=1)

    add_body(doc,
        "This delivery package includes the complete source code, sample data, configuration "
        "templates, and deployment-ready files for the DCF Valuation Engine. Below is what "
        "is included:"
    )

    add_heading_styled(doc, "Complete Deliverables", level=2)

    add_table(doc,
        ["Deliverable", "Description", "Format"],
        [
            ["Web Dashboard", "10-tab interactive analysis with live re-computation on every input change", "Browser-based (Flask)"],
            ["Excel Workbook", "17-tab formula-linked .xlsx — every cell references the Assumptions sheet", ".xlsx (openpyxl)"],
            ["PDF Investment Memo", "10-15 page formatted memo with all valuation outputs", ".pdf (ReportLab)"],
            ["Full Source Code", "Complete Python engine — modular, well-documented, extensible", "Python 3.10+"],
            ["Sample Data", "3-year historical financials (CSV) with working configuration", "CSV + JSON"],
            ["Deployment Files", "One-click Windows launcher, Vercel config, WSGI compatibility", "BAT / JSON"],
        ],
        col_widths=[1.5, 3.5, 1.6]
    )

    add_heading_styled(doc, "Technology Stack", level=2)

    add_table(doc,
        ["Layer", "Technology"],
        [
            ["Backend", "Python 3.10+, Flask 3.0"],
            ["Frontend", "React 18 (CDN), Chart.js 4, Babel JSX"],
            ["Excel Generation", "openpyxl (formula injection, not static values)"],
            ["PDF Generation", "ReportLab (optional)"],
            ["Data Processing", "pandas, numpy"],
            ["Live Market Data", "yfinance, SEC EDGAR XBRL API (optional)"],
            ["Deployment", "Local / Vercel serverless / any WSGI server"],
        ],
        col_widths=[1.8, 4.8]
    )

    add_body(doc,
        "No Node.js installation is required. The frontend uses CDN-hosted libraries and runs "
        "directly in the browser with zero build step."
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  3. UNIQUE SELLING POINTS
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Unique Selling Points (USP)", level=1)

    add_body(doc,
        "What sets this engine apart from typical DCF templates, generic financial modelling "
        "tools, or manual spreadsheets:"
    )

    usps = [
        ("Formula-Linked Excel (Not Value Dumps)",
         "Most automated models export static numbers. This engine injects live Excel formulas — "
         "Revenue, COGS, GP, EBITDA, D&A, Interest, Tax, FCF, DCF, Terminal Value — all reference "
         "the Assumptions sheet. Change one input in Excel and the entire 17-tab model recalculates "
         "natively. This provides full auditability and buyer confidence."),

        ("16-Step Integrated Pipeline",
         "The engine runs a complete 3-statement model (Income Statement → Working Capital → "
         "Capex & D&A → Multi-Tranche Debt → Balance Sheet → Cash Flow) before computing WACC, "
         "DCF, Scenarios, Monte Carlo, Sensitivity, Tornado, and Comps. Every step feeds into the "
         "next — depreciation from the Capex schedule overwrites the IS, interest from the Debt "
         "schedule flows through, and the Balance Sheet uses cash as a plug to force A = L + E."),

        ("Institutional-Quality Analytics",
         "10,000-iteration Monte Carlo simulation with 5 randomised drivers, 2D sensitivity "
         "heatmaps (WACC vs Terminal Growth, Revenue Growth vs EBITDA Margin), side-by-side "
         "Base/Bull/Bear scenario comparison, and tornado charts showing ranked driver impact "
         "on equity value."),

        ("Zero Build-Step Frontend",
         "The interactive dashboard is a single HTML file using React 18 and Chart.js via CDN. "
         "No Node.js, no npm, no webpack — just launch the Python server and open a browser. "
         "The dashboard auto-reruns the model within 800ms of any input change."),

        ("Dual-Delivery Model",
         "One set of inputs produces both real-time interactive analytics (for exploration and "
         "presentation) and a portable, formula-linked Excel workbook (for due diligence and "
         "offline analysis). The PDF memo adds a third output for deal documentation."),

        ("Fully Self-Contained & Deployable",
         "Runs locally with a single command, deploys to Vercel serverless with zero configuration, "
         "or scales on any WSGI server (Gunicorn, Waitress). No database — all state is in JSON "
         "configs and CSV data files."),
    ]

    for title, desc in usps:
        add_heading_styled(doc, title, level=3)
        add_body(doc, desc)

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  4. PLATFORM WALKTHROUGH
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Platform Walkthrough", level=1)

    add_heading_styled(doc, "How It Works — End to End", level=2)

    add_body(doc,
        "The platform follows a simple workflow from data input to deliverable output:"
    )

    steps = [
        ("Step 1: Prepare Input Data",
         "Place a CSV file with historical financials in the data/ folder. The engine accepts "
         "tidy format (period, account, amount, statement columns) or wide format. A sample "
         "dataset (Asian Street Eats, 2022–2024) is included."),
        ("Step 2: Configure Assumptions",
         "Create or edit a JSON config file with your forecast assumptions — revenue growth "
         "method, margin percentages, working capital days, capex approach, WACC inputs, "
         "terminal value parameters, debt tranches, and Monte Carlo settings. A fully worked "
         "example config is included."),
        ("Step 3: Launch the Dashboard",
         "Run 'python dashboard_api.py' (or double-click run_dashboard.bat on Windows). "
         "The server starts on http://localhost:5050 and opens 10 analysis tabs with all model "
         "outputs computed live."),
        ("Step 4: Adjust & Explore",
         "Use the collapsible sidebar to change any assumption. The model re-runs within 800ms "
         "and all charts, tables, and KPI cards update automatically. Load preset configurations "
         "from the dropdown."),
        ("Step 5: Export Deliverables",
         "Click 'Export Excel' in the sidebar to download the 17-tab formula-linked workbook. "
         "The PDF memo is also generated if ReportLab is installed."),
    ]
    for title, desc in steps:
        add_heading_styled(doc, title, level=3)
        add_body(doc, desc)

    add_heading_styled(doc, "Data Flow Architecture", level=2)

    add_body(doc,
        "The system operates as a standard request-response API. The frontend captures user "
        "assumptions in a collapsible sidebar, sends them as a JSON payload to the Python "
        "backend, which constructs a typed configuration and runs the 16-step pipeline. "
        "The pipeline returns structured data that the dashboard renders as charts and tables."
    )

    add_table(doc,
        ["Stage", "Action", "Output"],
        [
            ["1. User Input", "Sidebar controls adjusted", "React state update"],
            ["2. API Call", "POST /api/run (JSON payload)", "Config construction"],
            ["3. Pipeline", "16-step computation", "PipelineResult object"],
            ["4. Response", "JSON serialisation", "Charts, tables, KPIs"],
            ["5. Export", "POST /api/export-excel", "Formula-linked .xlsx"],
        ],
        col_widths=[1.5, 2.8, 2.3]
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  5. DASHBOARD — 10 TABS
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Dashboard — 10-Tab Breakdown", level=1)

    add_body(doc,
        "The interactive dashboard presents the complete valuation model across 10 purpose-built "
        "analysis tabs. Every tab updates live when any input changes."
    )

    dashboard_tabs = [
        ["Overview",
         "KPI summary cards (Revenue, EBITDA, Enterprise Value, Equity Value, Price per Share), "
         "revenue & margin trend charts, and an equity bridge waterfall."],
        ["Income Statement",
         "Projected P&L with waterfall chart, stacked revenue breakdown, full IS table with "
         "CAGR/YoY/manual growth, scenario-aware depreciation and interest."],
        ["Balance Sheet",
         "IFRS IAS 1 ordered assets vs liabilities composition chart, full BS table with "
         "cash plug to enforce Assets = Liabilities + Equity."],
        ["Cash Flow",
         "CFO (indirect method), CFI, and CFF breakdown chart. Unlevered and Levered FCF "
         "analysis with cross-statement reconciliation."],
        ["WACC",
         "CAPM build-up table (Rf + β×ERP + Size + Country Risk), synthetic credit rating "
         "(Damodaran grid), cost of debt, blended WACC calculation."],
        ["DCF",
         "Discount factor schedule with mid-year or end-of-period convention. Terminal value "
         "(Gordon Growth + Exit Multiple blend), enterprise-to-equity bridge."],
        ["Scenarios",
         "Side-by-side Base / Bull / Bear comparison with configurable revenue multipliers, "
         "margin shifts, terminal adjustments, and capex overrides."],
        ["Sensitivity",
         "Two 2D heatmap tables: WACC vs Terminal Growth Rate, and Revenue Growth vs "
         "EBITDA Margin. Colour-coded with the base case highlighted."],
        ["Monte Carlo",
         "10,000-iteration simulation histogram with 5 randomised drivers (Revenue Growth, "
         "EBITDA Margin, WACC, Terminal Growth, Exit Multiple). Statistics: Mean, Median, "
         "Std Dev, P10, P25, P75, P90, probability of exceeding base case."],
        ["Tornado",
         "Ranked bar chart showing ±20% driver swings and their impact on equity value. "
         "Identifies which assumptions have the greatest sensitivity."],
    ]

    for tab_info in dashboard_tabs:
        add_heading_styled(doc, tab_info[0], level=3)
        add_body(doc, tab_info[1])

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  6. EXCEL EXPORT
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Excel Export — 17-Tab Workbook", level=1)

    add_body(doc,
        "The Excel export is not a data dump. It is a fully formula-linked financial model. "
        "The engine writes native Excel formulas into every cell so the workbook is computationally "
        "alive after export. Change any value on the Assumptions sheet and the entire model "
        "recalculates natively in Excel."
    )

    add_heading_styled(doc, "Formula-Linked Tabs (Live Excel Formulas)", level=2)

    add_table(doc,
        ["Tab #", "Tab Name", "Content"],
        [
            ["1", "Cover", "Company name, model date, version"],
            ["2", "Assumptions", "All model inputs — the single source of truth"],
            ["3", "Income Statement", "Revenue → NI with formulas referencing Assumptions"],
            ["4", "Working Capital", "DSO/DIO/DPO-driven AR, Inventory, AP, NWC, ΔNWC"],
            ["5", "Capex & D&A", "PP&E rollforward with half-year convention"],
            ["6", "Debt Schedule", "Multi-tranche with amortisation, bullet maturity, cash sweep"],
            ["7", "Balance Sheet", "IFRS-ordered with cash plug (A = L + E enforced)"],
            ["8", "Cash Flow", "Indirect CFO + CFI + CFF, UFCF and LFCF"],
            ["9", "WACC", "CAPM cost of equity + synthetic credit rating cost of debt"],
            ["10", "DCF", "PV of FCFs, terminal value blend, equity bridge, price per share"],
            ["11", "Checks", "Balance sheet balance, cross-statement reconciliation"],
        ],
        col_widths=[0.5, 1.8, 4.3]
    )

    add_heading_styled(doc, "Analytics Tabs (Python-Computed Values)", level=2)

    add_table(doc,
        ["Tab #", "Tab Name", "Content"],
        [
            ["12", "Scenarios", "Base / Bull / Bear side-by-side comparison"],
            ["13", "Sensitivity", "2D heatmaps (WACC vs TG, Revenue Growth vs Margin)"],
            ["14", "Monte Carlo", "10,000 simulation rows with histogram and statistics"],
            ["15", "Tornado", "Driver sensitivity ranking with embedded chart"],
            ["16", "Comps", "Comparable company multiples (EV/Revenue, EV/EBITDA, P/E)"],
            ["17", "Audit Trail", "Pipeline execution log for model traceability"],
        ],
        col_widths=[0.5, 1.8, 4.3]
    )

    add_heading_styled(doc, "How Formula Injection Works", level=2)

    add_body(doc,
        "The engine uses a deterministic row/column mapping system. Each financial line item "
        "has a fixed row number, and all formulas reference the Assumptions sheet using absolute "
        "cell references (e.g., Assumptions!$B$4 for Revenue CAGR). This means:"
    )

    add_bullet(doc, "Changing Revenue CAGR in cell B4 of Assumptions recalculates all Revenue, GP, EBITDA, and downstream values.")
    add_bullet(doc, "Changing COGS % recalculates Gross Profit through to Net Income and FCF.")
    add_bullet(doc, "Changing WACC inputs recalculates discount factors, PV of FCFs, and equity value.")
    add_bullet(doc, "The Balance Sheet remains balanced because cash is a formula-driven plug.")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  7. FILE & FOLDER REFERENCE
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Complete File & Folder Reference", level=1)

    add_heading_styled(doc, "Root Directory", level=2)

    add_table(doc,
        ["File / Folder", "Purpose"],
        [
            ["dashboard_api.py", "Flask web server — routes, config parsing, JSON serialisation, export"],
            ["run_dashboard.bat", "Windows one-click launcher (double-click to start)"],
            ["requirements.txt", "Python dependencies"],
            ["config.example.json", "Template configuration with default assumptions"],
            ["config.asian_street_ib_grade.json", "Working example config (Food & Beverage company)"],
            ["vercel.json", "Vercel serverless deployment configuration"],
            ["templates/dashboard.html", "Single-file React 18 dashboard (entire frontend)"],
            ["data/", "Historical financial data (CSV files)"],
            ["api/index.py", "Vercel serverless entry point"],
        ],
        col_widths=[2.8, 3.8]
    )

    add_heading_styled(doc, "Core Engine — src/dcf_engine/", level=2)

    add_table(doc,
        ["File", "Purpose"],
        [
            ["config.py", "Typed configuration dataclasses with JSON serialisation and alias handling"],
            ["pipeline.py", "16-step orchestrator — the heart of the engine"],
            ["main.py", "CLI entry point (python -m src.dcf_engine.main)"],
        ],
        col_widths=[2.5, 4.1]
    )

    add_heading_styled(doc, "Financial Statements — src/dcf_engine/statements/", level=2)

    add_table(doc,
        ["File", "Purpose"],
        [
            ["income_statement.py", "Projects IS with CAGR / YoY / manual revenue, scenario overrides"],
            ["balance_sheet.py", "Constructs BS from IS + schedules; cash is plug to force A=L+E"],
            ["cash_flow.py", "CFO (indirect) + CFI + CFF; computes UFCF and LFCF"],
        ],
        col_widths=[2.5, 4.1]
    )

    add_heading_styled(doc, "Supporting Schedules — src/dcf_engine/schedules/", level=2)

    add_table(doc,
        ["File", "Purpose"],
        [
            ["working_capital.py", "DSO/DIO/DPO-driven WC schedule with NWC and Delta NWC"],
            ["capex_depreciation.py", "PP&E rollforward with half-year convention for new capex"],
            ["debt_schedule.py", "Multi-tranche debt with amortisation, bullet maturity, cash sweeps"],
        ],
        col_widths=[2.5, 4.1]
    )

    add_heading_styled(doc, "Valuation & Analytics — src/dcf_engine/valuation/", level=2)

    add_table(doc,
        ["File", "Purpose"],
        [
            ["wacc.py", "CAPM cost of equity + synthetic credit rating for cost of debt"],
            ["dcf_engine.py", "Core DCF: discount FCFs, terminal value (Gordon + Exit), equity bridge"],
            ["scenarios.py", "Side-by-side scenario comparison builder"],
            ["monte_carlo.py", "10,000-iteration MC simulation with 5 randomised drivers"],
            ["sensitivity.py", "2D sensitivity tables (WACC/TG and RevGrowth/Margin)"],
            ["tornado.py", "±20% driver swing analysis for tornado charts"],
        ],
        col_widths=[2.5, 4.1]
    )

    add_heading_styled(doc, "Export Layer — src/dcf_engine/output/", level=2)

    add_table(doc,
        ["File", "Purpose"],
        [
            ["excel_builder.py", "Orchestrates workbook creation (calls all sheet builders)"],
            ["excel_formats.py", "Shared formatting: fonts, fills, borders, Assumptions row map"],
            ["sheets_core.py", "Formula-linked sheets: Cover, Assumptions, IS, WC, Capex DA, Debt"],
            ["sheets_valuation.py", "Formula-linked sheets: Balance Sheet, Cash Flow, WACC, DCF"],
            ["sheets_analytics.py", "Value-based tabs: Scenarios, Sensitivity, MC, Tornado, Comps"],
            ["pdf_memo.py", "ReportLab PDF investment memo (graceful fallback if not installed)"],
        ],
        col_widths=[2.5, 4.1]
    )

    add_heading_styled(doc, "Other Modules", level=2)

    add_table(doc,
        ["File / Folder", "Purpose"],
        [
            ["comps/comps.py", "Comparable companies via yfinance (EV/Revenue, EV/EBITDA, P/E)"],
            ["ingestion/file_loader.py", "Flexible CSV/XLSX loader with auto-format detection"],
            ["ingestion/edgar_client.py", "SEC EDGAR XBRL API client for public company filings"],
            ["ingestion/market_data.py", "Live data helpers via yfinance (beta, Rf, ERP)"],
        ],
        col_widths=[2.5, 4.1]
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  8. INSTALLATION & SETUP
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Installation & Setup Guide", level=1)

    add_heading_styled(doc, "Prerequisites", level=2)

    add_table(doc,
        ["Requirement", "Version", "Notes"],
        [
            ["Python", "3.10 or higher", "Download from python.org"],
            ["pip", "Latest", "Included with Python"],
            ["Git", "Any", "For version control (optional)"],
        ],
        col_widths=[1.5, 1.5, 3.6]
    )

    add_body(doc, "No Node.js installation is required.", space_before=4)

    add_heading_styled(doc, "Installation Steps", level=2)

    install_steps = [
        "Create a virtual environment (recommended):\n    python -m venv .venv",
        "Activate the virtual environment:\n    Windows:  .venv\\Scripts\\activate\n    macOS/Linux:  source .venv/bin/activate",
        "Install core dependencies:\n    pip install -r requirements.txt",
        "Install optional dependencies for full functionality:\n    pip install matplotlib scipy reportlab yfinance",
    ]
    for i, step in enumerate(install_steps, 1):
        p = doc.add_paragraph()
        p.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(f"{i}. {step}")
        set_font(run, size=8.5, color=COLOR_DARK, name="Consolas")

    add_heading_styled(doc, "Running the Dashboard", level=2)

    add_table(doc,
        ["Method", "Command", "Notes"],
        [
            ["Direct Python", "python dashboard_api.py", "Opens on http://localhost:5050"],
            ["Windows Batch", "Double-click run_dashboard.bat", "Same as above, one-click"],
            ["Custom Port", "flask --app dashboard_api run --port 8080", "For port configuration"],
        ],
        col_widths=[1.5, 3.0, 2.1]
    )

    add_heading_styled(doc, "Optional Dependencies", level=2)

    add_table(doc,
        ["Package", "Purpose", "Required?"],
        [
            ["yfinance", "Live beta, risk-free rate, comparable company data", "No"],
            ["reportlab", "PDF investment memo generation", "No"],
            ["matplotlib", "Additional chart rendering in exports", "No"],
            ["scipy", "Statistical distributions for Monte Carlo", "No"],
        ],
        col_widths=[1.5, 3.5, 1.6]
    )

    add_body(doc,
        "The core engine (dashboard + Excel export) runs with only the packages in requirements.txt. "
        "Optional packages add PDF generation, live market data, and enhanced statistics."
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  9. LOADING YOUR OWN DATA
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. How to Load Your Own Data", level=1)

    add_heading_styled(doc, "Step 1: Prepare Your CSV", level=2)

    add_body(doc,
        "Place a CSV file in the data/ folder. The engine accepts tidy format with "
        "four columns: period, account, amount, statement."
    )

    add_body(doc, "Example CSV format:", space_before=4)

    csv_example = [
        ["period", "account", "amount", "statement"],
        ["2024", "Revenue", "5000000", "income"],
        ["2024", "COGS", "2250000", "income"],
        ["2024", "SGA", "1000000", "income"],
        ["2024", "Cash", "500000", "balance"],
        ["2024", "PP&E", "800000", "balance"],
    ]
    table = doc.add_table(rows=len(csv_example), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(csv_example):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.space_before = Pt(2)
            p.space_after = Pt(2)
            run = p.add_run(val)
            if r_idx == 0:
                set_font(run, size=8, bold=True, color=COLOR_WHITE, name="Consolas")
                set_cell_shading(cell, "1B3A5C")
            else:
                set_font(run, size=8, color=COLOR_DARK, name="Consolas")
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "F7F9FB")
    doc.add_paragraph().space_after = Pt(4)

    add_heading_styled(doc, "Step 2: Create a Configuration File", level=2)

    add_body(doc,
        "Copy config.example.json and customise the assumptions for your company. "
        "Set the 'data_file' field to point to your CSV (e.g., \"data/my_company.csv\"). "
        "Adjust revenue growth, margins, working capital days, capex method, WACC inputs, "
        "and valuation parameters."
    )

    add_heading_styled(doc, "Step 3: Load in the Dashboard", level=2)

    add_body(doc,
        "Start the dashboard, then use the preset dropdown to load your config file. "
        "Alternatively, run from the command line:"
    )

    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run("    python -m src.dcf_engine.main --config config.your_company.json")
    set_font(run, size=8.5, color=COLOR_DARK, name="Consolas")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  10. CONFIGURATION REFERENCE
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Configuration Reference", level=1)

    add_body(doc,
        "All model assumptions are stored in a single JSON configuration file. "
        "Below are the main sections and their key parameters."
    )

    add_heading_styled(doc, "Company Section", level=2)
    add_table(doc,
        ["Parameter", "Type", "Description"],
        [
            ["name", "string", "Company name displayed in outputs"],
            ["ticker", "string", "Stock ticker (optional, for live data pulls)"],
            ["industry", "string", "Industry classification"],
            ["currency", "string", "Reporting currency (e.g., USD, GBP)"],
        ],
        col_widths=[1.8, 1.0, 3.8]
    )

    add_heading_styled(doc, "Forecast Section", level=2)
    add_table(doc,
        ["Parameter", "Type", "Description"],
        [
            ["years", "integer", "Number of projection years (typically 5)"],
            ["revenue_method", "string", "Growth method: cagr, yoy, or manual"],
            ["revenue_cagr", "decimal", "Compound annual growth rate (e.g., 0.08 for 8%)"],
            ["cogs_pct_revenue", "decimal", "COGS as percentage of revenue"],
            ["opex_pct_revenue", "decimal", "SGA & other operating expenses as % of revenue"],
            ["tax_rate", "decimal", "Effective tax rate"],
            ["dso / dio / dpo", "integer", "Working capital days (Sales Outstanding, Inventory, Payable)"],
            ["capex_method", "string", "pct_revenue, fixed, or manual"],
            ["depreciation_rate", "decimal", "Annual depreciation rate on PP&E"],
        ],
        col_widths=[1.8, 1.0, 3.8]
    )

    add_heading_styled(doc, "WACC Section", level=2)
    add_table(doc,
        ["Parameter", "Type", "Description"],
        [
            ["risk_free_rate", "decimal", "Government bond yield (e.g., 0.043 for 4.3%)"],
            ["market_risk_premium", "decimal", "Equity risk premium (e.g., 0.055 for 5.5%)"],
            ["beta", "decimal", "Equity beta (levered)"],
            ["size_premium", "decimal", "Small-cap premium (optional)"],
            ["country_risk_premium", "decimal", "Country-specific risk (optional)"],
            ["target_debt_weight", "decimal", "Target debt / total capital"],
            ["target_equity_weight", "decimal", "Target equity / total capital"],
            ["interest_coverage_ratio", "decimal", "For synthetic credit rating / cost of debt"],
        ],
        col_widths=[2.2, 1.0, 3.4]
    )

    add_heading_styled(doc, "Valuation Section", level=2)
    add_table(doc,
        ["Parameter", "Type", "Description"],
        [
            ["terminal_growth_rate", "decimal", "Long-term growth rate for Gordon Growth model"],
            ["exit_ev_ebitda_multiple", "decimal", "Exit EV/EBITDA multiple for terminal value"],
            ["discount_convention", "string", "mid_year or end_period"],
            ["gordon_weight", "decimal", "Blend weight for Gordon vs Exit Multiple (0.0–1.0)"],
            ["cash / debt", "number", "Net cash and debt for equity bridge"],
            ["fully_diluted_shares", "integer", "For per-share equity value calculation"],
        ],
        col_widths=[2.4, 1.0, 3.2]
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  11. DEPLOYMENT OPTIONS
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. Deployment Options", level=1)

    add_table(doc,
        ["Method", "Command / Action", "Notes"],
        [
            ["Local (Development)", "python dashboard_api.py", "Runs on http://localhost:5050"],
            ["Windows One-Click", "Double-click run_dashboard.bat", "Installs deps + launches"],
            ["Vercel (Serverless)", "vercel --prod", "Pre-configured, zero setup"],
            ["Gunicorn (Linux)", "gunicorn dashboard_api:app --bind 0.0.0.0:8080", "Production WSGI"],
            ["Waitress (Windows)", "waitress-serve --port=8080 dashboard_api:app", "Production WSGI"],
        ],
        col_widths=[1.8, 3.2, 1.6]
    )

    add_body(doc,
        "The app has no database dependency. All state is in JSON config files and CSV data files, "
        "making deployment trivial on any platform that supports Python."
    )

    add_heading_styled(doc, "Vercel Deployment Details", level=2)

    add_body(doc,
        "The project includes vercel.json which routes all requests to api/index.py. "
        "On Vercel, Monte Carlo iterations are automatically capped at 2,000 and live market "
        "data pulls are disabled for serverless compatibility. Install the Vercel CLI with "
        "'npm i -g vercel', then run 'vercel --prod' from the project root."
    )

    add_heading_styled(doc, "Environment Variables", level=2)

    add_table(doc,
        ["Variable", "Required", "Purpose"],
        [
            ["VERCEL", "No", "Auto-detected on Vercel; enables serverless mode"],
            ["VERCEL_ENV", "No", "Same as VERCEL"],
        ],
        col_widths=[2.0, 1.0, 3.6]
    )

    add_body(doc, "No .env file is needed for local development.")

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  12. EXTENSION & CUSTOMISATION GUIDE
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. Extension & Customisation Guide", level=1)

    add_body(doc,
        "The codebase is modular and extensible. Below are common customisation scenarios "
        "with step-by-step instructions."
    )

    add_heading_styled(doc, "Adding a New Valuation Model (e.g., LBO)", level=2)

    lbo_steps = [
        "Create a new module under src/dcf_engine/valuation/ (e.g., lbo.py) with assumptions, debt waterfall logic, and return metrics (IRR/MOIC).",
        "Extend src/dcf_engine/config.py with typed LBO configuration sections.",
        "Integrate as a new pipeline step in src/dcf_engine/pipeline.py without breaking the existing DCF path.",
        "Expose output fields in dashboard_api.py for frontend rendering.",
        "Add an export tab writer in the output/ directory for the LBO Excel sheet.",
    ]
    for step in lbo_steps:
        add_bullet(doc, step)

    add_heading_styled(doc, "Updating Excel Formatting or Adding Rows", level=2)

    excel_steps = [
        "Update style tokens and formats in src/dcf_engine/output/excel_formats.py.",
        "If adding rows, update row-constant maps and all dependent formula references.",
        "Keep the Assumptions row map (the 'A' dictionary) synchronised to avoid broken references.",
        "Validate by exporting the workbook and changing assumptions directly in Excel to confirm full recalculation.",
    ]
    for step in excel_steps:
        add_bullet(doc, step)

    add_heading_styled(doc, "Adding a New Dashboard Tab", level=2)

    tab_steps = [
        "Add computation logic as a new pipeline step in pipeline.py.",
        "Include the output in the JSON response from dashboard_api.py.",
        "Add a new tab component in templates/dashboard.html (React JSX).",
        "Wire Chart.js or table rendering to the new data fields.",
    ]
    for step in tab_steps:
        add_bullet(doc, step)

    add_heading_styled(doc, "Design Rule", level=3)
    add_body(doc,
        "Always preserve the separation between the calculation layer (Python objects in "
        "valuation/ and statements/) and the presentation/export layer (API serialisation "
        "and Excel sheet writers in output/). This keeps the engine testable and portable."
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  13. API REFERENCE
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "13. API Reference", level=1)

    add_body(doc,
        "All backend endpoints are served by Flask from dashboard_api.py."
    )

    add_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/", "GET", "Serve the dashboard HTML page"],
            ["/api/run", "POST", "Run full pipeline, return all model results as JSON"],
            ["/api/export-excel", "POST", "Run pipeline, return formula-linked .xlsx download"],
            ["/api/configs", "GET", "List available preset config files"],
            ["/api/config/<filename>", "GET", "Load a specific config JSON"],
            ["/api/data-files", "GET", "List CSV/XLSX files in data/ folder"],
            ["/api/data-file-base-values/<file>", "GET", "Extract base-year values from historical CSV"],
        ],
        col_widths=[2.8, 0.8, 3.0]
    )

    add_heading_styled(doc, "POST /api/run Payload Structure", level=2)

    add_body(doc,
        "The JSON payload sent to /api/run contains these top-level sections:"
    )

    add_table(doc,
        ["Section", "Description"],
        [
            ["company", "Name, ticker, industry, currency"],
            ["forecast", "Projection years, revenue method/rates, margins, capex, WC, tax"],
            ["wacc", "CAPM inputs: Rf, ERP, beta, size/country premium, capital weights"],
            ["valuation", "Terminal growth, exit multiple, blend weight, bridge items, shares"],
            ["monte_carlo", "Distribution means and standard deviations for 5 drivers"],
            ["sensitivity", "Range parameters for 2D heatmap tables"],
            ["scenarios", "Base/Bull/Bear override multipliers"],
            ["debt_tranches", "Array of debt tranche objects (balance, rate, amort, maturity)"],
            ["data_file", "Path to historical CSV (optional)"],
        ],
        col_widths=[1.8, 4.8]
    )

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  14. VERSION HISTORY
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "14. Version History & Quality Assurance", level=1)

    add_body(doc,
        "The engine has been through 13 major versions with systematic bug fixing, feature "
        "additions, and cross-check testing. Below is a summary of key releases."
    )

    add_table(doc,
        ["Version", "Key Changes"],
        [
            ["V13.0", "Current release"],
            ["V12.0", "Fixed Monte Carlo Excel NORM.INV #NAME? errors; fixed histogram column widths"],
            ["V10.0", "Dashboard typography refresh (Verdana/Consolas); Excel BS total label fix; client packaging"],
            ["V9.0", "Auto-sync Monte Carlo means from base case; frontend useEffect auto-sync"],
            ["V5.0", "Auto-compute base retained earnings; BS/CF cash reconciliation fix"],
            ["V4.0", "CF→BS ordering fix; shares in millions fix; credit spread surfacing; amortisation cap"],
            ["V3.0", "Monte Carlo EBITDA margin fix; scenario IS re-linking; TV blend weight; mid-year convention"],
        ],
        col_widths=[1.0, 5.6]
    )

    add_heading_styled(doc, "Quality Assurance", level=2)

    add_body(doc,
        "The engine includes comprehensive test suites that validate:"
    )

    qa_items = [
        "Income Statement projections (CAGR, YoY, manual methods)",
        "Working Capital schedule (DSO/DIO/DPO-driven)",
        "Capex & Depreciation with half-year convention",
        "Multi-tranche debt schedule with amortisation and cash sweeps",
        "IS re-linking (D&A from Capex, Interest from Debt schedules)",
        "Balance Sheet balancing (Assets = Liabilities + Equity, cash as plug)",
        "Cash Flow to Balance Sheet cash reconciliation",
        "Auto retained earnings computation",
        "WACC calculation (CAPM + synthetic rating)",
        "DCF valuation (mid-year convention, terminal value blend)",
        "UFCF from IS/CF cross-check",
        "Excel formula-to-Python value matching ($0.00 precision)",
        "Scenario IS with full re-linking",
        "Tornado driver sensitivity",
        "Monte Carlo distribution parameter sync",
    ]
    for item in qa_items:
        add_bullet(doc, item)

    add_page_break(doc)

    # ═════════════════════════════════════════════════════════════════
    #  15. SUPPORT & NEXT STEPS
    # ═════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "15. Support & Next Steps", level=1)

    add_heading_styled(doc, "Getting Started Checklist", level=2)

    checklist = [
        "Install Python 3.10+ and pip",
        "Run: pip install -r requirements.txt",
        "Run: python dashboard_api.py",
        "Open http://localhost:5050 in your browser",
        "Load a preset config from the dropdown (Asian Street Eats is included)",
        "Explore all 10 dashboard tabs",
        "Click 'Export Excel' to download the 17-tab formula-linked workbook",
        "Open the workbook, go to the Assumptions tab, change a value, and watch formulas recalculate",
        "When ready, load your own CSV data and create a custom config JSON",
    ]
    for i, item in enumerate(checklist, 1):
        p = doc.add_paragraph()
        p.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(f"  {i}.  {item}")
        set_font(run, size=8.5, color=COLOR_DARK)

    add_heading_styled(doc, "Recommended Enhancements (Optional)", level=2)

    add_table(doc,
        ["Enhancement", "Description"],
        [
            ["Excel YoY/Manual Revenue", "Extend IS formulas for non-CAGR revenue methods"],
            ["Historical Tab in Excel", "Add historical data sheet with VLOOKUP references"],
            ["Multi-Currency Support", "Currency conversion for international models"],
            ["LBO Module", "Leveraged buyout model with IRR/MOIC analysis"],
            ["Docker Container", "One-command deployment via containerisation"],
        ],
        col_widths=[2.2, 4.4]
    )

    add_heading_styled(doc, "Contact", level=2)

    add_body(doc,
        "For questions, customisation requests, or technical support, please contact "
        "the development team."
    )

    # ─── Footer note ─────────────────────────────────────────────────
    doc.add_paragraph().space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 50)
    set_font(run, size=10, color=COLOR_ACCENT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(4)
    run2 = p2.add_run("IB-Grade DCF Valuation Engine  ·  Version 13.0  ·  Confidential")
    set_font(run2, size=8, color=COLOR_MEDIUM, italic=True)

    # ─── Save ────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Document saved to: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
